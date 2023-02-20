from docx import Document
from docx.shared import RGBColor

# Function to highlight exact matches between two documents
def highlight_exact_matches(main_doc, compare_doc, output_doc):
    # Open the main document
    main_doc = Document(main_doc)
    # Open the comparison document
    compare_doc = Document(compare_doc)
    # Create a new document for the output
    output_doc = Document()

    # Iterate through the paragraphs in the main document
    for main_p in main_doc.paragraphs:
        # Create a new paragraph for the output
        output_p = output_doc.add_paragraph()
        # Flag to indicate whether we found an exact match in the comparison document
        found_exact_match = False
        # Iterate through the paragraphs in the comparison document
        for compare_p in compare_doc.paragraphs:
            # If the paragraphs are an exact match, add the main paragraph to the output document and highlight it in green
            if main_p.text == compare_p.text:
                for run in main_p.runs:
                    output_run = output_p.add_run(run.text)
                    if run.bold:
                        output_run.bold = True
                    if run.italic:
                        output_run.italic = True
                    if run.underline:
                        output_run.underline = True
                    if run.font.highlight_color:
                        output_run.font.highlight_color = run.font.highlight_color
                    output_run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)  # Green
                found_exact_match = True
                break

        # If we didn't find an exact match in the comparison document, copy the main paragraph to the output document without highlighting
        if not found_exact_match:
            for run in main_p.runs:
                output_run = output_p.add_run(run.text)
                if run.bold:
                    output_run.bold = True
                if run.italic:
                    output_run.italic = True
                if run.underline:
                    output_run.underline = True
                if run.font.highlight_color:
                    output_run.font.highlight_color = run.font.highlight_color
                output_run.font.color.rgb = run.font.color.rgb

        # Set the output paragraph style to match the main paragraph style
        output_p.style = main_p.style

    # Save the output document
    output_doc.save(output_doc_path)

# Path to the main document
main_doc_path = '/Users/aj/Desktop/Year 4/Spring Semester/ICSI 499/Milestone 3/Testing1.docx'
# Path to the comparison document
compare_doc_path = '/Users/aj/Desktop/Year 4/Spring Semester/ICSI 499/Milestone 3/Testing2.docx'
# Path to the output document
output_doc_path = '/Users/aj/Desktop/Year 4/Spring Semester/ICSI 499/Milestone 3/output.docx'

# Highlight exact matches between the two documents and save the output
highlight_exact_matches(main_doc_path, compare_doc_path, output_doc_path)
